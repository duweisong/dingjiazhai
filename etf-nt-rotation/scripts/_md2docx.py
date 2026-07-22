"""Convert a Markdown file to .docx using python-docx. Usage: python _md2docx.py <input.md>"""
import re, sys
from docx import Document
from docx.shared import Pt

def md2docx(src: str, dst: str):
    doc = Document()
    with open(src, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # code blocks
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue

        # headings
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        # tables
        elif line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|') and '---' in lines[i+1]:
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                if '---' in tl and not tl.strip('|- '):
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        table.cell(ri, ci).text = cell_text
            continue
        # hr / empty / image / comment
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif not line.strip():
            pass
        elif line.startswith('![') or line.startswith('<!--'):
            i += 1
            continue
        else:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            text = re.sub(r'^> ', '', text)
            doc.add_paragraph(text)
        i += 1

    doc.save(dst)
    print(f'DOCX saved: {dst}')

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python _md2docx.py <input.md> <output.docx>')
        sys.exit(1)
    md2docx(sys.argv[1], sys.argv[2])
